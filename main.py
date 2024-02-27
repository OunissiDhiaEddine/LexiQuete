from docx import Document
import os 
from prettytable import PrettyTable
#for needed imports and testing check readme.txt



""" hadi function that count how many words in a file , also it depends on a language selection """

def count_words(file_path, language='english'):
    try:
        if file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                if language == 'arabic':
                    word_count = len(text.split())
                else:
                    word_count = len(text.split())
                return word_count
        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            if language == 'arabic':
                word_count = sum(len(paragraph.text.split()) for paragraph in doc.paragraphs)
            else:
                word_count = sum(len(paragraph.text.split()) for paragraph in doc.paragraphs)
            return word_count
        else:
            return "Unsupported file format. Please provide a .txt or .docx file."
    except FileNotFoundError:
        return "File not found."
    



  
""" hadi function removes stopwords (depending on selected language) """

def remove_stopwords(text, language='english'):
    if language == 'arabic':
        stopwords = set(['في', 'على', 'من', 'إلى', 'هو', 'أن', 'ب', 'و', 'يكون', 'فيه', 'كان', 'هذا', 'هذه', 'ذلك', 'هناك', 'الذي', 'هذه', 'هؤلاء', 'ذلكم', 'تلك', 'هذين', 'هذان', 'هاتان', 'هاتين', 'هؤلاء', 'هذه'
])  
    else:
        stopwords = set(['a', 'an', 'the', 'in', 'on', 'at', 'and', 'or', 'but', 'for', 'to', 'of', 'with', 'by', 'as', 'from', 'into', 'onto', 'over', 'under', 'among', 'between', 'within', 'without', 'through', 'during', 'before', 'after', 'since', 'until', 'about', 'against', 'across', 'along', 'around', 'off', 'out', 'up', 'down', 'through'
])  
    cleaned_text = ' '.join(word for word in text.split() if word.lower() not in stopwords)
    return cleaned_text


""" hadi function dir table that contains frequency and percentage of appearance of words in the cleared text and order 
 them from top to bot """

def show_word_frequency_table(cleaned_text):
    words = cleaned_text.split()
    word_count = len(words)
    word_frequency = {}
    
    # Count word frequency
    for word in words:
        word_frequency[word] = word_frequency.get(word, 0) + 1
    
    # Calculate percentage of appearance
    word_percentage = {word: (count / word_count) * 100 for word, count in word_frequency.items()}
    
    # Create and populate the table
    table = PrettyTable(['Word', 'Frequency', 'Percentage'])
    sorted_word_frequency = sorted(word_frequency.items(), key=lambda x: x[1], reverse=True)
    for word, frequency in sorted_word_frequency:
        percentage = word_percentage[word]
        table.add_row([word, frequency, f'{percentage:.2f}%'])
    
    # Print the table
    print(table)




""" main program that rn now : asks for path and language , counts words , saves a edited txt without stopwords in same folder
  , shows a frequency and percentage in a table  """

def main():
    file_path = input("Enter the path to the text file or docx file: ")
    language = input("Select language (arabic/english): ")
    word_count = count_words(file_path, language.lower())
    if isinstance(word_count, int):
        print(f"The file contains {word_count} words.")
        with open(file_path, 'r', encoding='utf-8') as file:
            original_text = file.read()
            cleaned_text = remove_stopwords(original_text, language.lower())
            new_file_path = os.path.splitext(file_path)[0] + "_cleared" + os.path.splitext(file_path)[1]
            with open(new_file_path, 'w', encoding='utf-8') as new_file:
                new_file.write(cleaned_text)
        print(f"Cleaned text saved to {new_file_path}.")
    else:
        print(word_count)

    show_word_frequency_table(cleaned_text)

if __name__ == "__main__":
    main()